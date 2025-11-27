from flask import Flask, request, send_file, render_template_string, jsonify
import PyPDF2
from docx import Document
import os
import uuid
import glob
from datetime import datetime, timedelta
import time
import logging

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # 5 MB
UPLOAD_FOLDER = "uploads"
MAX_FILE_SIZE = 5 * 1024 * 1024

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Sadə logging konfiqurasiyası
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def cleanup_old_files():
    """1 saatdan köhnə faylları təmizlə"""
    try:
        cutoff_time = datetime.now() - timedelta(hours=1)
        files_removed = 0
        for file_path in glob.glob(os.path.join(UPLOAD_FOLDER, "*")):
            if os.path.isfile(file_path):
                file_time = datetime.fromtimestamp(os.path.getctime(file_path))
                if file_time < cutoff_time:
                    try:
                        os.remove(file_path)
                        files_removed += 1
                        logger.info(f"Köhnə fayl silindi: {file_path}")
                    except Exception as e:
                        logger.error(f"Fayl silinə bilmədi {file_path}: {e}")
        if files_removed > 0:
            logger.info(f"Ümumi {files_removed} köhnə fayl silindi")
    except Exception as e:
        logger.error(f"Fayl təmizləmə xətası: {e}")

def convert_pdf_to_docx(pdf_path, docx_path):
    """PDF-dən DOCX-ə çevir"""
    try:
        doc = Document()
        
        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            total_pages = len(pdf_reader.pages)
            
            for page_num in range(total_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text and text.strip():
                        # Mətni təmizlə
                        text = clean_text(text)
                        
                        # Əgər ilk səhifə deyilsə, səhifə əlavə et
                        if page_num > 0:
                            doc.add_page_break()
                        
                        # Mətni əlavə et
                        paragraph = doc.add_paragraph(text)
                        
                        logger.info(f"Səhifə {page_num + 1}/{total_pages} çevrildi")
                    else:
                        logger.warning(f"Səhifə {page_num + 1} boş və ya mətn tapılmadı")
                        if page_num > 0:
                            doc.add_page_break()
                        doc.add_paragraph(f"[Səhifə {page_num + 1} - Mətn tapılmadı]")
                        
                except Exception as e:
                    logger.error(f"Səhifə {page_num + 1} xətası: {e}")
                    if page_num > 0:
                        doc.add_page_break()
                    doc.add_paragraph(f"[Səhifə {page_num + 1} çevrilə bilmədi]")
                    continue
        
        # DOCX faylını yadda saxla
        doc.save(docx_path)
        logger.info(f"DOCX faylı yaradıldı: {docx_path}")
        return True
        
    except Exception as e:
        logger.error(f"Çevirmə xətası: {e}")
        return False

def clean_text(text):
    """Mətni təmizlə və formatla"""
    if not text:
        return ""
    
    # XML-uyğun olmayan simvolları təmizlə
    cleaned = ""
    for char in text:
        if ord(char) >= 32 or char in '\n\t\r':
            cleaned += char
    
    # Artıq boşluqları təmizlə
    lines = cleaned.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line = line.strip()
        if line:
            # Sətirdəki artıq boşluqları təmizlə
            words = line.split()
            cleaned_line = ' '.join(words)
            cleaned_lines.append(cleaned_line)
    
    return '\n'.join(cleaned_lines)

HTML = """
<!DOCTYPE html>
<html>
<head>
    <title>PDF → Word | AxtarGet</title>
    <meta charset="utf-8">
    <script src="https://cdn.tailwindcss.com"></script>
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
</head>
<body class="bg-gradient-to-br from-purple-900 to-black min-h-screen flex items-center justify-center p-4">
    <div class="bg-white/10 backdrop-blur-lg rounded-3xl p-6 md:p-10 max-w-2xl w-full shadow-2xl border border-white/20">
        <h1 class="text-3xl md:text-5xl font-black text-center bg-gradient-to-r from-cyan-400 to-purple-400 bg-clip-text text-transparent mb-4">
            PDF → Word Çevirici
        </h1>
        <p class="text-gray-300 text-center mb-6 md:mb-8 text-sm md:text-base">Bütün səhifələr tam şəkildə çevrilir</p>
        
        <form method="post" enctype="multipart/form-data" class="space-y-6" id="upload-form">
            <div class="border-2 border-dashed border-cyan-400 rounded-2xl p-6 md:p-10 text-center hover:border-cyan-300 transition" id="drop-zone">
                <input type="file" name="pdf" accept=".pdf" required class="hidden" id="file-input">
                <label for="file-input" class="cursor-pointer block h-full">
                    <div class="text-4xl md:text-6xl mb-3 md:mb-4">📄</div>
                    <p class="text-lg md:text-xl text-cyan-300 font-bold" id="file-text">PDF faylı seç və ya bura sürükle</p>
                    <p class="text-xs md:text-sm text-gray-400 mt-2">Maksimum ölçü: 5 MB</p>
                </label>
            </div>
            
            <div id="file-info" class="hidden p-4 bg-cyan-500/20 border border-cyan-400 rounded-xl">
                <div class="flex items-center justify-between">
                    <div>
                        <p class="text-cyan-300 font-bold text-sm md:text-lg">Seçilmiş fayl:</p>
                        <p id="filename" class="text-white text-xs md:text-sm mt-1"></p>
                        <p id="filesize" class="text-gray-300 text-xs mt-1"></p>
                    </div>
                    <button type="button" onclick="clearFile()" class="text-red-400 hover:text-red-300 text-xl">×</button>
                </div>
            </div>

            <div id="progress-container" class="hidden">
                <div class="w-full bg-gray-700 rounded-full h-3 overflow-hidden mb-2">
                    <div id="progress-fill" class="bg-gradient-to-r from-cyan-500 to-purple-600 h-full w-0 transition-all duration-500 ease-out"></div>
                </div>
                <p id="progress-text" class="text-cyan-300 text-sm text-center">Hazırlanır...</p>
            </div>
            
            <button type="submit" class="w-full py-4 md:py-6 bg-gradient-to-r from-cyan-500 to-purple-600 text-white text-xl md:text-2xl font-black rounded-2xl hover:scale-105 transition transform duration-200 disabled:opacity-50 disabled:cursor-not-allowed" id="submit-btn">
                WORD-Ə ÇEVİR
            </button>
        </form>
        
        <div id="result-container"></div>
        
        <div class="mt-6 text-center">
            <button onclick="resetForm()" class="inline-block px-6 py-3 bg-gray-600 text-white font-bold rounded-xl hover:bg-gray-700 transition">
                🗑️ Formanı Təmizlə
            </button>
        </div>
        
        <p class="text-center text-gray-500 mt-8 md:mt-10 text-xs md:text-sm">© 2025 AxtarGet</p>
    </div>

    <script>
        let currentFile = null;
        const MAX_FILE_SIZE = 5 * 1024 * 1024;

        // Elementləri seç
        const dropZone = document.getElementById('drop-zone');
        const fileInput = document.getElementById('file-input');
        const fileText = document.getElementById('file-text');
        const fileInfo = document.getElementById('file-info');
        const fileName = document.getElementById('filename');
        const fileSize = document.getElementById('filesize');
        const submitBtn = document.getElementById('submit-btn');
        const progressContainer = document.getElementById('progress-container');
        const progressFill = document.getElementById('progress-fill');
        const progressText = document.getElementById('progress-text');
        const resultContainer = document.getElementById('result-container');
        const uploadForm = document.getElementById('upload-form');

        // Fayl seçildikdə
        fileInput.addEventListener('change', function(e) {
            if (this.files && this.files[0]) {
                handleFileSelect(this.files[0]);
            }
        });

        function handleFileSelect(file) {
            if (file.size > MAX_FILE_SIZE) {
                showError('Fayl çox böyükdür! Maksimum 5 MB qəbul edilir.');
                resetFileInput();
                return;
            }
            
            if (file.type !== 'application/pdf') {
                showError('Yalnız PDF faylları qəbul edilir!');
                resetFileInput();
                return;
            }
            
            currentFile = file;
            updateFileInfo(file);
        }

        function updateFileInfo(file) {
            fileName.textContent = file.name;
            fileSize.textContent = formatFileSize(file.size);
            fileInfo.classList.remove('hidden');
            fileText.textContent = 'Fayl seçildi';
            dropZone.classList.add('border-green-400', 'bg-green-500/10');
            dropZone.classList.remove('border-cyan-400');
            submitBtn.disabled = false;
        }

        function clearFile() {
            resetFileInput();
            fileInfo.classList.add('hidden');
            fileText.textContent = 'PDF faylı seç və ya bura sürükle';
            dropZone.classList.remove('border-green-400', 'bg-green-500/10');
            dropZone.classList.add('border-cyan-400');
            submitBtn.disabled = true;
            currentFile = null;
        }

        function resetFileInput() {
            fileInput.value = '';
            currentFile = null;
        }

        function resetForm() {
            clearFile();
            resultContainer.innerHTML = '';
            progressContainer.classList.add('hidden');
            submitBtn.disabled = false;
        }

        function formatFileSize(bytes) {
            if (bytes === 0) return '0 Bytes';
            const k = 1024;
            const sizes = ['Bytes', 'KB', 'MB'];
            const i = Math.floor(Math.log(bytes) / Math.log(k));
            return parseFloat((bytes / Math.pow(k, i)).toFixed(2)) + ' ' + sizes[i];
        }

        function showError(message) {
            resultContainer.innerHTML = `
                <div class="mt-6 p-4 bg-red-500/20 border border-red-400 rounded-2xl text-center">
                    <p class="text-red-300 font-bold">❌ ${message}</p>
                </div>
            `;
        }

        function showSuccess(message, downloadUrl) {
            resultContainer.innerHTML = `
                <div class="mt-6 p-4 bg-green-500/20 border border-green-400 rounded-2xl text-center animate-pulse">
                    <p class="text-green-300 font-bold text-lg mb-3">${message}</p>
                    <a href="${downloadUrl}" class="inline-block px-6 py-3 bg-green-600 text-white font-bold rounded-xl hover:bg-green-700 transition">
                        📥 WORD FAYLINI ENDİR
                    </a>
                    <p class="text-gray-300 text-sm mt-2">Yeni fayl çevirmək üçün yuxarıdan başqa PDF seçin</p>
                </div>
            `;
        }

        function updateProgress(percent, text) {
            progressFill.style.width = percent + '%';
            progressText.textContent = text;
        }

        // Drag & Drop funksionallığı
        ['dragenter', 'dragover', 'dragleave', 'drop'].forEach(eventName => {
            dropZone.addEventListener(eventName, preventDefaults, false);
            document.body.addEventListener(eventName, preventDefaults, false);
        });

        function preventDefaults(e) {
            e.preventDefault();
            e.stopPropagation();
        }

        ['dragenter', 'dragover'].forEach(eventName => {
            dropZone.addEventListener(eventName, highlight, false);
        });

        ['dragleave', 'drop'].forEach(eventName => {
            dropZone.addEventListener(eventName, unhighlight, false);
        });

        function highlight() {
            dropZone.classList.add('border-green-400', 'bg-cyan-500/20');
        }

        function unhighlight() {
            if (!currentFile) {
                dropZone.classList.remove('border-green-400', 'bg-cyan-500/20');
            }
        }

        dropZone.addEventListener('drop', handleDrop, false);

        function handleDrop(e) {
            const dt = e.dataTransfer;
            const files = dt.files;
            
            if (files.length > 0) {
                handleFileSelect(files[0]);
            }
        }

        // Form göndərilməsi
        uploadForm.addEventListener('submit', async function(e) {
            e.preventDefault();
            
            if (!currentFile) {
                showError('Zəhmət olmasa bir PDF faylı seçin!');
                return;
            }

            const formData = new FormData();
            formData.append('pdf', currentFile);

            // UI-ni yenilə
            submitBtn.disabled = true;
            progressContainer.classList.remove('hidden');
            resultContainer.innerHTML = '';
            
            updateProgress(10, 'Fayl yüklənir...');

            try {
                updateProgress(30, 'PDF oxunur...');
                
                const response = await fetch('/', {
                    method: 'POST',
                    body: formData
                });

                updateProgress(70, 'Word sənədi hazırlanır...');

                const html = await response.text();
                
                // HTML cavabını parse et
                const parser = new DOMParser();
                const doc = parser.parseFromString(html, 'text/html');
                
                const errorElem = doc.querySelector('.bg-red-500\\\\/20');
                const successElem = doc.querySelector('.bg-green-500\\\\/20');
                
                if (errorElem) {
                    const errorText = errorElem.textContent.trim();
                    showError(errorText);
                } else if (successElem) {
                    const successText = successElem.querySelector('.text-green-300').textContent.trim();
                    const downloadLink = successElem.querySelector('a');
                    const downloadUrl = downloadLink ? downloadLink.href : null;
                    
                    updateProgress(100, 'Tamamlandı!');
                    showSuccess(successText, downloadUrl);
                } else {
                    showError('Gözlənilməz xəta baş verdi!');
                }

            } catch (error) {
                showError('Şəbəkə xətası: ' + error.message);
            } finally {
                setTimeout(() => {
                    progressContainer.classList.add('hidden');
                    submitBtn.disabled = false;
                }, 2000);
            }
        });
    </script>
</body>
</html>
"""

@app.route("/", methods=["GET", "POST"])
def index():
    # Hər dəfə səhifə yüklənəndə köhnə faylları təmizlə
    cleanup_old_files()
    
    if request.method == "POST":
        if 'pdf' not in request.files:
            return render_template_string(HTML, error="❌ Fayl seçilməyib!")
        
        pdf_file = request.files['pdf']
        
        if pdf_file.filename == '':
            return render_template_string(HTML, error="❌ Fayl seçilməyib!")
        
        if not pdf_file.filename.lower().endswith('.pdf'):
            return render_template_string(HTML, error="❌ Yalnız PDF faylı qəbul edilir!")
        
        # Unikal fayl adları yarat
        unique_id = str(uuid.uuid4())
        pdf_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}.pdf")
        docx_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}.docx")
        
        try:
            # PDF faylını yadda saxla
            pdf_file.save(pdf_path)
            
            # Fayl ölçüsünü yoxla
            file_size = os.path.getsize(pdf_path)
            if file_size > MAX_FILE_SIZE:
                os.remove(pdf_path)
                return render_template_string(HTML, error="❌ Fayl çox böyükdür! Maksimum 5 MB.")
            
            logger.info(f"PDF yükləndi: {pdf_path} ({file_size} bytes)")
            
            # Çevirmə prosesi
            update_progress = lambda p, t: None  # Progress funksiyası
            
            success = convert_pdf_to_docx(pdf_path, docx_path)
            
            # PDF faylını sil
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
            
            if success and os.path.exists(docx_path):
                return render_template_string(
                    HTML, 
                    result="✅ PDF uğurla Word sənədinə çevrildi!", 
                    filename=f"{unique_id}.docx"
                )
            else:
                return render_template_string(
                    HTML, 
                    error="❌ PDF çevrilə bilmədi. Zəhmət olmasa başqa fayl sınayın."
                )
                
        except Exception as e:
            logger.error(f"Ümumi xəta: {e}")
            # Təmizlik
            for path in [pdf_path, docx_path]:
                if os.path.exists(path):
                    try:
                        os.remove(path)
                    except:
                        pass
            
            return render_template_string(
                HTML, 
                error=f"❌ Xəta: {str(e)}"
            )
    
    return render_template_string(HTML)

@app.route("/download/<filename>")
def download(filename):
    try:
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        if os.path.exists(file_path) and filename.endswith('.docx'):
            return send_file(
                file_path,
                as_attachment=True,
                download_name="cevirilmis_word.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        return "Fayl tapılmadı", 404
    except Exception as e:
        logger.error(f"Download xətası: {e}")
        return "Xəta baş verdi", 500

@app.route("/health")
def health():
    return "OK"

@app.route("/cleanup")
def cleanup_route():
    cleanup_old_files()
    return "Təmizlik tamamlandı"

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    # Production üçün debug False
    app.run(host="0.0.0.0", port=port, debug=False)
